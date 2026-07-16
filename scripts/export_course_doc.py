#!/usr/bin/env python3
"""Export the full Tutor-Based AI Expert course to a portable Word workbook."""

from __future__ import annotations

import re
import subprocess
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
EXPORTS = ROOT / "exports"
WEEKS_DIR = ROOT / "curriculum" / "weeks"

DOC_FILES = [
    ("Course Overview", ROOT / "README.md"),
    ("Curriculum Index", ROOT / "curriculum" / "INDEX.md"),
    ("Resource Library", ROOT / "curriculum" / "RESOURCES.md"),
    ("Knowledge Checks & Phase Gates", ROOT / "curriculum" / "CHECKS.md"),
    ("Phase Gate Checklist", ROOT / "checkpoints" / "phase-gates.md"),
    ("Capstone Project Guide", ROOT / "capstone" / "README.md"),
]


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    new_run.append(r_pr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_notes_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFBFBF")
        borders.append(element)
    tbl_pr.append(borders)


def add_notes_area(doc: Document, heading: str, rows: int = 7) -> None:
    p = doc.add_paragraph()
    run = p.add_run(heading)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    table = doc.add_table(rows=rows, cols=1)
    table.autofit = True
    set_notes_table_borders(table)
    for row in table.rows:
        cell = row.cells[0]
        cell.text = ""
        for para in cell.paragraphs:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            para.paragraph_format.space_after = Pt(6)
            para.add_run(" ").font.size = Pt(11)
    doc.add_paragraph()


def add_inline_markdown(paragraph, text: str) -> None:
    pattern = re.compile(r"(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*|[^*\[]+)")
    for match in pattern.finditer(text):
        if match.group(2) and match.group(3):
            add_hyperlink(paragraph, match.group(2), match.group(3))
        elif match.group(4):
            run = paragraph.add_run(match.group(4))
            run.bold = True
        else:
            paragraph.add_run(match.group(0))


def render_markdown(doc: Document, content: str, stop_at_day: bool = False) -> list[str]:
    """Render markdown; return day section headers found (for inserting notes)."""
    lines = content.splitlines()
    i = 0
    day_headers: list[str] = []
    in_code = False

    while i < len(lines):
        line = lines[i]

        if line.strip().startswith("```"):
            in_code = not in_code
            i += 1
            continue

        if in_code:
            p = doc.add_paragraph(line)
            p.style = "Intense Quote"
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue

        if line.startswith("### "):
            title = line[4:].strip()
            doc.add_heading(title, level=3)
            if title.startswith("Day "):
                day_headers.append(title)
            i += 1
            continue

        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
            i += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph("—" * 40)
            i += 1
            continue

        if line.startswith("|") and "|" in line[1:]:
            table_rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                row_line = lines[i].strip()
                if re.match(r"^\|[-:\s|]+\|$", row_line):
                    i += 1
                    continue
                cells = [c.strip() for c in row_line.strip("|").split("|")]
                table_rows.append(cells)
                i += 1
            if table_rows:
                cols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(table_rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < cols:
                            table.rows[r_idx].cells[c_idx].text = cell_text
            continue

        bullet = re.match(r"^(\s*)([-*]|\d+\.)\s+(.*)$", line)
        if bullet:
            indent, _, body = bullet.groups()
            style = "List Bullet" if bullet.group(2) in "-*" else "List Number"
            p = doc.add_paragraph(style=style)
            if len(indent) >= 2:
                p.paragraph_format.left_indent = Inches(0.25 * (len(indent) // 2))
            add_inline_markdown(p, body)
            i += 1
            continue

        if line.strip().startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, line.strip()[2:])
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, line)
        i += 1

    return day_headers


def render_week_with_notes(doc: Document, week_path: Path) -> None:
    content = week_path.read_text(encoding="utf-8")
    sections = re.split(r"(?=^### Day \d+)", content, flags=re.MULTILINE)
    preamble = sections[0]
    day_sections = sections[1:]

    render_markdown(doc, preamble)

    for day_block in day_sections:
        day_match = re.match(r"^### (Day \d+[^\n]*)", day_block)
        day_title = day_match.group(1) if day_match else "Day"
        render_markdown(doc, day_block)
        add_notes_area(doc, f"📝 My Notes — {day_title}", rows=8)

    add_notes_area(doc, "📝 Weekly Reflection — key takeaways, blockers, next week", rows=10)


def build_title_page(doc: Document) -> None:
    for _ in range(6):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = 1
    run = title.add_run("Tutor-Based AI Expert Course")
    run.bold = True
    run.font.size = Pt(28)

    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    sub_run = subtitle.add_run("Complete Workbook · 78 Weeks · 390 Hours")
    sub_run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = 1
    meta.add_run(f"Exported {date.today().isoformat()}\n").font.size = Pt(11)
    meta.add_run("60 minutes per business day · Mon–Fri").font.size = Pt(11)

    doc.add_paragraph()
    guide = doc.add_paragraph()
    guide.alignment = 1
    g = guide.add_run(
        "Portable study edition with embedded lessons, documentation, "
        "and note spaces for every day."
    )
    g.italic = True
    g.font.size = Pt(11)

    add_page_break(doc)


def build_how_to_use(doc: Document) -> None:
    doc.add_heading("How to Use This Workbook", level=1)
    items = [
        "Work one day block per business day (60 minutes). Each day lists embedded source links — open them from this document.",
        "After each day lesson, write in the My Notes box: summaries, questions, code snippets, and deliverable status.",
        "At the end of each week, complete the Weekly Reflection box before moving on.",
        "On phase gate weeks (8, 16, 24, 32, 40, 52, 78), complete the Knowledge Checks section and log your score in notes.",
        "Build artifacts in this repository under capstone/ and notes/.",
        "Regenerate this export anytime: python3 scripts/export_course_doc.py",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Number")
    add_page_break(doc)


def build_toc(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    doc.add_paragraph("Course Documentation", style="List Bullet")
    for title, _ in DOC_FILES:
        doc.add_paragraph(f"  {title}", style="List Bullet")

    phases = [
        ("Phase 1 — AI & RAG Foundations", 1, 8),
        ("Phase 2 — Knowledge Maps", 9, 16),
        ("Phase 3 — Graph DB & Student Modeling", 17, 24),
        ("Phase 4 — GraphRAG", 25, 32),
        ("Phase 5 — MCP & Agents", 33, 40),
        ("Phase 6 — Capstone & Production", 41, 52),
        ("Phase 7 — Expert Mastery", 53, 78),
    ]
    doc.add_paragraph("Weekly Lessons (with daily note spaces)", style="List Bullet")
    for phase_name, start, end in phases:
        doc.add_paragraph(f"  {phase_name}: Weeks {start:02d}–{end:02d}", style="List Bullet")
    add_page_break(doc)


def build_documentation(doc: Document) -> None:
    doc.add_heading("Course Documentation", level=1)
    for title, path in DOC_FILES:
        if not path.exists():
            continue
        doc.add_heading(title, level=2)
        render_markdown(doc, path.read_text(encoding="utf-8"))
        add_notes_area(doc, f"📝 Notes — {title}", rows=5)
        add_page_break(doc)


def build_lessons(doc: Document) -> None:
    doc.add_heading("Weekly Lessons", level=1)
    doc.add_paragraph(
        "Each week includes five 60-minute day blocks. Write in the note boxes "
        "after every day and complete the weekly reflection on Friday."
    )

    week_files = sorted(WEEKS_DIR.glob("week-*.md"))
    for week_path in week_files:
        doc.add_heading(week_path.stem.replace("-", " ").title(), level=2)
        render_week_with_notes(doc, week_path)
        add_page_break(doc)


def convert_to_doc(docx_path: Path, doc_path: Path) -> bool:
    try:
        subprocess.run(
            ["textutil", "-convert", "doc", str(docx_path), "-output", str(doc_path)],
            check=True,
            capture_output=True,
            text=True,
        )
        return doc_path.exists()
    except (subprocess.CalledProcessError, FileNotFoundError) as exc:
        print(f"Warning: could not convert to .doc: {exc}", file=sys.stderr)
        return False


def main() -> None:
    EXPORTS.mkdir(parents=True, exist_ok=True)
    docx_path = EXPORTS / "tutor-ai-expert-course.docx"
    doc_path = EXPORTS / "tutor-ai-expert-course.doc"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    build_title_page(doc)
    build_how_to_use(doc)
    build_toc(doc)
    build_documentation(doc)
    build_lessons(doc)

    doc.save(docx_path)
    print(f"Wrote {docx_path} ({docx_path.stat().st_size // 1024} KB)")

    if convert_to_doc(docx_path, doc_path):
        print(f"Wrote {doc_path} ({doc_path.stat().st_size // 1024} KB)")
    else:
        print("Legacy .doc conversion skipped — use the .docx file (opens in Word).")


if __name__ == "__main__":
    main()
